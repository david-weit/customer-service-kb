"""DOCX 解析器（仿 RAGFlow RAGFlowDocxParser）。"""

from io import BytesIO
from pathlib import Path
from typing import BinaryIO, List, Tuple, Union

from docx import Document as DocxDocument


class DocxParser:
    """用 python-docx 抽取段落与表格，表格行转为自然语言键值句。"""

    def __extract_table_content(self, table) -> List[str]:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            return [" | ".join(c for c in row if c) for row in rows if any(row)]

        headers = rows[0]
        lines = []
        for row in rows[1:]:
            cells = []
            for i, value in enumerate(row):
                if not value:
                    continue
                header = headers[i] if i < len(headers) and headers[i] else ""
                cells.append(f"{header}: {value}" if header else value)
            if cells:
                lines.append("; ".join(cells))
        return lines

    def __call__(
        self,
        fnm: Union[str, Path, bytes, BinaryIO],
        from_page: int = 0,
        to_page: int = 100000,
    ) -> Tuple[List[Tuple[str, str]], List[List[str]]]:
        """
        Returns:
            secs: [(text, style_name), ...]
            tbls: [table_lines, ...] 每个 table_lines 为行文本列表
        """
        if isinstance(fnm, (str, Path)):
            doc = DocxDocument(str(fnm))
        elif isinstance(fnm, bytes):
            doc = DocxDocument(BytesIO(fnm))
        else:
            doc = DocxDocument(fnm)

        pn = 0
        secs: List[Tuple[str, str]] = []
        for paragraph in doc.paragraphs:
            if pn > to_page:
                break
            text_parts = []
            for run in paragraph.runs:
                if pn > to_page:
                    break
                if from_page <= pn < to_page and paragraph.text.strip():
                    text_parts.append(run.text)
                if "lastRenderedPageBreak" in run._element.xml:
                    pn += 1
            text = "".join(text_parts).strip()
            if text:
                style = (
                    paragraph.style.name
                    if paragraph.style and hasattr(paragraph.style, "name")
                    else ""
                )
                secs.append((text, style))

        tbls = [self.__extract_table_content(tb) for tb in doc.tables]
        return secs, tbls
